from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


##هنا قمنا بتخزين الشفرة
filename = 'output.txt'
myString = ''
with open(filename, 'r') as f:
    myString = f.readline()

##هنا قمنا بتقسيم الشفرة وتحويلها الى قائمة
myList = myString.split('-')

##هنا قمنا بانشاء دالة لتحويل النص الى قائمة
def listConverter(str):
    return str.split('.')

meaning ={
    'PR': "مدير إدارة المخاطر",
    'VIC1': "مشرف الوحدة الخامسة",
}

##هنا قمنا بالتحقق من الرقم الاول في القائمة
if myList[0] == '1':
    ##هنا قمنا بانشاء ملف وكتابة الرسالة
    document = Document()

    ##هنا قمنا بتحديد العنوان واضفنا بعض النصوص المطلوبة
    paragraph = document.add_paragraph('')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('سعادة ' +meaning.get(myList[1])+'  في الجامعة                                 حفظه الله')
    run.bold = True
    run.font.size = Pt(18)
    run = paragraph.add_run('\n السلام عليكم ورحمة الله وبركاته\n\n')
    run.bold = True
    run.font.size = Pt(18)

    paragraph = document.add_paragraph("")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('مرفق لسعادتكم أسماء فريق لجنة الطوارئ في وحدتنا')
    run.font.size = Pt(18)

    ##هنا قمنا بالتحقق من النص الرابع في القائمة
    if myList[4][0] == 'T':
        rows = int(myList[4][1])
        columns = int(myList[4][3])
        ##هنا قمنا بانشاء جدول وتحديد عدد الصفوف والاعمدة
        table = document.add_table(rows=rows, cols=columns)
        tableHeader = myList[5].split('.')

        ##هنا قمنا بتحديد العناوين للجدول
        for i in range(rows):
            item = listConverter(myList[i+5])
            for j in range(columns):
                cell = table.cell(i, j)
                cell.text = item[j]

                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


                for run in paragraph.runs:
                    run.font.size = Pt(18)
                    if i == 0:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="cccccc"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '4')
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)

                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '4')
                left.set(qn('w:space'), '0')
                left.set(qn('w:color'), 'auto')
                tcBorders.append(left)

                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'auto')
                tcBorders.append(bottom)

                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:sz'), '4')
                right.set(qn('w:space'), '0')
                right.set(qn('w:color'), 'auto')
                tcBorders.append(right)

                tcPr.append(tcBorders)


    space = "\u00A0"
    paragraph = document.add_paragraph('')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("يرجى الاطلاع والتوجيه"+space*5+"\n\n")
    run.font.size = Pt(18)

    paragraph = document.add_paragraph('')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("وتقبلوا تحياتي"+space*21+"\n\n")
    run.font.size = Pt(18)
    run.bold = True

    paragraph = document.add_paragraph('')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(meaning.get(myList[2])+space*72 )
    run.font.size = Pt(18)  
    run.bold = True

    document.save('test.docx')

